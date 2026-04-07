"""
Document processing actions for reading and extracting text from Word/PDF documents.
"""
from __future__ import annotations

import os
from pathlib import Path
from typing import Optional, Dict, Any, List


def read_pdf_text(file_path: str, max_pages: Optional[int] = None) -> str:
    """
    Extract text from a PDF file.

    Args:
        file_path: Path to the PDF file.
        max_pages: Maximum number of pages to read (None for all pages).

    Returns:
        Extracted text content from the PDF.

    Raises:
        FileNotFoundError: If the PDF file does not exist.
        ValueError: If the file is not a valid PDF.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"PDF file not found: {file_path}")

    if not file_path.lower().endswith('.pdf'):
        raise ValueError(f"File is not a PDF: {file_path}")

    try:
        import pypdf
    except ImportError:
        try:
            import PyPDF2 as pypdf
        except ImportError:
            raise ImportError("pypdf or PyPDF2 is required for PDF processing. Install with: pip install pypdf")

    try:
        with open(file_path, 'rb') as f:
            reader = pypdf.PdfReader(f)
            total_pages = len(reader.pages)

            if max_pages is not None:
                total_pages = min(total_pages, max_pages)

            text_parts = []
            for i in range(total_pages):
                page = reader.pages[i]
                text_parts.append(page.extract_text())

            return '\n\n'.join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to read PDF: {e}")


def read_word_text(file_path: str) -> str:
    """
    Extract text from a Microsoft Word document (.docx).

    Args:
        file_path: Path to the Word document.

    Returns:
        Extracted text content from the Word document.

    Raises:
        FileNotFoundError: If the Word file does not exist.
        ValueError: If the file is not a valid Word document.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Word document not found: {file_path}")

    if not file_path.lower().endswith(('.docx', '.docm')):
        raise ValueError(f"File is not a Word document: {file_path}")

    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for Word processing. Install with: pip install python-docx")

    try:
        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs]
        return '\n\n'.join(paragraphs)
    except Exception as e:
        raise ValueError(f"Failed to read Word document: {e}")


def get_document_info(file_path: str) -> Dict[str, Any]:
    """
    Get metadata and information about a document.

    Args:
        file_path: Path to the document file.

    Returns:
        Dictionary containing document metadata.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file type is not supported.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Document file not found: {file_path}")

    info: Dict[str, Any] = {
        'file_name': os.path.basename(file_path),
        'file_size': os.path.getsize(file_path),
        'file_type': Path(file_path).suffix.lower(),
    }

    if file_path.lower().endswith('.pdf'):
        try:
            import pypdf
        except ImportError:
            try:
                import PyPDF2 as pypdf
            except ImportError:
                return info

        with open(file_path, 'rb') as f:
            reader = pypdf.PdfReader(f)
            info['page_count'] = len(reader.pages)
            if reader.metadata:
                info['title'] = str(reader.metadata.get('/Title', ''))
                info['author'] = str(reader.metadata.get('/Author', ''))
                info['subject'] = str(reader.metadata.get('/Subject', ''))
                info['creator'] = str(reader.metadata.get('/Creator', ''))

    elif file_path.lower().endswith(('.docx', '.docm')):
        try:
            from docx import Document
        except ImportError:
            return info

        doc = Document(file_path)
        info['paragraph_count'] = len(doc.paragraphs)
        info['table_count'] = len(doc.tables)

    return info


def read_document_pages(
    file_path: str,
    start_page: int = 1,
    end_page: Optional[int] = None
) -> List[str]:
    """
    Read specific pages from a PDF document.

    Args:
        file_path: Path to the PDF file.
        start_page: Starting page number (1-indexed).
        end_page: Ending page number (None for start_page only).

    Returns:
        List of text strings, one per page.

    Raises:
        FileNotFoundError: If the PDF file does not exist.
        ValueError: If the file is not a valid PDF.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"PDF file not found: {file_path}")

    if not file_path.lower().endswith('.pdf'):
        raise ValueError(f"File is not a PDF: {file_path}")

    try:
        import pypdf
    except ImportError:
        try:
            import PyPDF2 as pypdf
        except ImportError:
            raise ImportError("pypdf or PyPDF2 is required for PDF processing.")

    try:
        with open(file_path, 'rb') as f:
            reader = pypdf.PdfReader(f)
            total_pages = len(reader.pages)

            if start_page < 1:
                start_page = 1
            if start_page > total_pages:
                return []

            end_page_idx = end_page if end_page is not None else start_page
            end_page_idx = min(end_page_idx, total_pages)

            pages = []
            for i in range(start_page - 1, end_page_idx):
                page = reader.pages[i]
                pages.append(page.extract_text())

            return pages
    except Exception as e:
        raise ValueError(f"Failed to read PDF pages: {e}")


def extract_document_words(file_path: str) -> List[str]:
    """
    Extract all words from a document as a list.

    Args:
        file_path: Path to the document file (PDF or Word).

    Returns:
        List of all words in the document.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file type is not supported.
    """
    import re

    if file_path.lower().endswith('.pdf'):
        text = read_pdf_text(file_path)
    elif file_path.lower().endswith(('.docx', '.docm')):
        text = read_word_text(file_path)
    else:
        raise ValueError(f"Unsupported document type: {file_path}")

    words = re.findall(r'\b\w+\b', text)
    return words


def search_document(file_path: str, query: str, case_sensitive: bool = False) -> List[Dict[str, Any]]:
    """
    Search for text within a document.

    Args:
        file_path: Path to the document file (PDF or Word).
        query: Text string to search for.
        case_sensitive: Whether the search should be case-sensitive.

    Returns:
        List of match dictionaries with 'text', 'page' (PDF) or 'paragraph' (Word).

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file type is not supported.
    """
    matches: List[Dict[str, Any]] = []

    if file_path.lower().endswith('.pdf'):
        try:
            import pypdf
        except ImportError:
            try:
                import PyPDF2 as pypdf
            except ImportError:
                raise ImportError("pypdf or PyPDF2 is required for PDF processing.")

        with open(file_path, 'rb') as f:
            reader = pypdf.PdfReader(f)
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                search_text = text if case_sensitive else text.lower()
                search_query = query if case_sensitive else query.lower()

                if search_query in search_text:
                    for line in text.split('\n'):
                        compare_line = line if case_sensitive else line.lower()
                        if search_query in compare_line:
                            matches.append({
                                'text': line.strip(),
                                'page': page_num
                            })

    elif file_path.lower().endswith(('.docx', '.docm')):
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required for Word processing.")

        doc = Document(file_path)
        for para_num, para in enumerate(doc.paragraphs, 1):
            text = para.text
            search_text = text if case_sensitive else text.lower()
            search_query = query if case_sensitive else query.lower()

            if search_query in search_text:
                matches.append({
                    'text': text.strip(),
                    'paragraph': para_num
                })

    else:
        raise ValueError(f"Unsupported document type: {file_path}")

    return matches


def count_document_words(file_path: str) -> int:
    """
    Count total words in a document.

    Args:
        file_path: Path to the document file.

    Returns:
        Total word count in the document.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file type is not supported.
    """
    return len(extract_document_words(file_path))


def get_document_headings(file_path: str) -> List[Dict[str, Any]]:
    """
    Extract headings from a Word document.

    Args:
        file_path: Path to the Word document.

    Returns:
        List of heading dictionaries with 'text', 'level', and 'paragraph_index'.

    Raises:
        FileNotFoundError: If the Word file does not exist.
        ValueError: If the file is not a Word document.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Word document not found: {file_path}")

    if not file_path.lower().endswith(('.docx', '.docm')):
        raise ValueError(f"File is not a Word document: {file_path}")

    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for Word processing.")

    doc = Document(file_path)
    headings: List[Dict[str, Any]] = []

    for idx, para in enumerate(doc.paragraphs):
        if para.style.name.startswith('Heading'):
            try:
                level = int(para.style.name.replace('Heading ', '').replace('Heading', '1'))
            except ValueError:
                level = 1

            headings.append({
                'text': para.text.strip(),
                'level': level,
                'paragraph_index': idx
            })

    return headings


# Aliases for compatibility
extract_text_from_pdf = read_pdf_text
extract_text_from_word = read_word_text
